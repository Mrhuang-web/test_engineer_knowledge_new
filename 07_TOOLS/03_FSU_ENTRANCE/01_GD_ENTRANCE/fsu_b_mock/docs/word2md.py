#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word to Markdown Converter
重点处理Word表格到Markdown表格的转换
"""

from docx import Document
from docx.shared import Pt
import re

def convert_word_to_markdown(word_file_path, md_file_path=None):
    """
    将Word文件转换为Markdown格式
    :param word_file_path: Word文件路径
    :param md_file_path: 输出Markdown文件路径，默认为原文件名.md
    :return: Markdown内容
    """
    # 加载Word文档
    doc = Document(word_file_path)
    
    markdown_content = []
    
    # 获取文档的所有块级元素（段落和表格）按照实际顺序
    blocks = []
    
    # 遍历文档的主体内容
    body = doc._element.body
    
    # 创建段落元素到段落对象的映射
    para_map = {}
    for i, para in enumerate(doc.paragraphs):
        para_map[para._element] = para
    
    # 创建表格元素到表格对象的映射
    table_map = {}
    for i, table in enumerate(doc.tables):
        table_map[table._element] = table
    
    for child in body:
        if child.tag.endswith('p'):
            # 这是一个段落
            if child in para_map:
                para = para_map[child]
                blocks.append(('para', para))
        elif child.tag.endswith('tbl'):
            # 这是一个表格
            if child in table_map:
                table = table_map[child]
                blocks.append(('table', table))
    
    # 按照实际顺序处理所有块级元素
    for block_type, block in blocks:
        if block_type == 'para':
            para = block
            # 处理标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1])
                markdown_content.append(f"{'#' * level} {para.text}\n")
            
            # 处理列表
            elif para.style.name in ['List Paragraph', 'List Bullet', 'List Number']:
                # 检查是否为有序列表
                if para.style.name == 'List Number':
                    # 简单处理，实际需要更复杂的逻辑来处理嵌套列表和正确编号
                    markdown_content.append(f"1. {para.text}\n")
                else:
                    # 无序列表
                    markdown_content.append(f"- {para.text}\n")
            
            # 处理普通文本
            elif para.text.strip():
                markdown_content.append(f"{para.text}\n")
            
            # 处理空行
            else:
                markdown_content.append("\n")
        elif block_type == 'table':
            table = block
            markdown_content.append("\n")  # 表格前添加空行
            
            # 转换表格
            table_md = convert_table_to_markdown(table)
            markdown_content.append(table_md)
            
            markdown_content.append("\n")  # 表格后添加空行
    
    # 合并所有内容
    markdown_text = ''.join(markdown_content)
    
    # 保存到文件
    if md_file_path:
        with open(md_file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
    
    return markdown_text

def convert_table_to_markdown(table):
    """
    将Word表格转换为Markdown表格
    :param table: docx表格对象
    :return: Markdown表格字符串
    """
    md_rows = []
    
    # 获取表格的行数和列数
    rows_count = len(table.rows)
    cols_count = len(table.columns)
    
    for i, row in enumerate(table.rows):
        cells = row.cells
        
        # 处理每个单元格的内容
        md_cells = []
        for cell in cells:
            # 合并单元格中的所有段落文本，使用换行符连接
            cell_text = '\n'.join(para.text.strip() for para in cell.paragraphs if para.text.strip())
            # 处理特殊字符，如|需要转义
            cell_text = cell_text.replace('|', '\\|')
            md_cells.append(cell_text)
        
        # 拼接当前行的Markdown格式
        md_row = '| ' + ' | '.join(md_cells) + ' |'
        md_rows.append(md_row)
        
        # 在表头下方添加分隔线
        if i == 0:
            # 为每列创建分隔线，默认居中对齐
            separators = ['---' for _ in range(cols_count)]
            md_separator = '| ' + ' | '.join(separators) + ' |'
            md_rows.append(md_separator)
    
    # 合并所有行
    return '\n'.join(md_rows)

def main():
    """
    命令行界面
    """
    import argparse
    
    parser = argparse.ArgumentParser(description='Word to Markdown Converter')
    parser.add_argument('input', help='Input Word file path')
    parser.add_argument('-o', '--output', help='Output Markdown file path', default=None)
    
    args = parser.parse_args()
    
    # 如果没有指定输出文件，使用输入文件名的.md扩展名
    if not args.output:
        args.output = args.input.rsplit('.', 1)[0] + '.md'
    
    # 执行转换
    print(f"Converting {args.input} to {args.output}...")
    convert_word_to_markdown(args.input, args.output)
    print("Conversion completed!")

if __name__ == '__main__':
    main()
